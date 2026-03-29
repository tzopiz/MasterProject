#!/usr/bin/env python3
"""
Создаёт DOCX с планом доклада на конференцию.
Текст: Times New Roman 14 pt, заголовки: 16 pt, интервал 1,5.
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    raise SystemExit("Установите: pip install python-docx")

def set_run_font(run, size_pt: int, bold: bool = False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold

def add_heading_custom(doc, text: str, level: int = 1):
    """Заголовок: Times New Roman 16 pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
    return p

def add_paragraph_custom(doc, text: str, bold_first: bool = False):
    """Абзац: Times New Roman 14 pt, интервал 1,5."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, bold=bold_first)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text: str, bold_label: bool = False):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, 14, bold=bold_label)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(3)
    return p

def main():
    doc = Document()
    # Стиль по умолчанию для всего документа
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    add_heading_custom(doc, "План доклада на конференцию: начальный этап проекта AI Doctor — анализ ВНЧС", 1)
    add_paragraph_custom(
        doc,
        "Объём: ~5 листов A4, шрифт Times New Roman 14 pt (текст), 16 pt (заголовки), "
        "межстрочный интервал 1,5 (ориентировочно 12–14 тыс. знаков с пробелами)."
    )
    add_paragraph_custom(
        doc,
        "Фокус: первые шаги, постановка задачи, что уже сделано. "
        "Проект в процессе; доклад — о начальном этапе."
    )

    add_heading_custom(doc, "1. Введение и актуальность (≈0,7–1 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Тема: интеллектуальная система для автоматизированного анализа височно-нижнечелюстного сустава (ВНЧС) по данным конусно-лучевой компьютерной томографии (КЛКТ) с использованием методов глубокого обучения.")
    add_bullet(doc, "Актуальность: заболевания ВНЧС широко распространены; КЛКТ — основной метод визуализации; ручной анализ 3D-снимков трудоёмок и субъективен; ИИ даёт возможность автоматизации и стандартизации.")
    add_bullet(doc, "Контекст: магистерская программа «Искусственный интеллект в медицине»; предшествующие работы — 2D U-Net для сегментации на сагиттальных срезах (ограничения: ручной выбор срезов, только 2D, малый датасет, отсутствие интеграции в приложение).")
    add_bullet(doc, "Цель доклада: представить постановку задачи и результаты начального этапа разработки системы.")
    add_paragraph_custom(doc, "Источники в проекте: Nir/otchet_nir_magistratura_sem1.md (Введение, актуальность), тезисы_конференция.txt.")

    add_heading_custom(doc, "2. Постановка задачи (≈1–1,2 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Общая задача: система автоматизированной оценки состояния ВНЧС по данным КЛКТ. Вход: DICOM-исследование (серия 2D-срезов). Выход: диагностическое заключение по левому и правому суставу (норма / патология; в перспективе — расширение классов).")
    add_bullet(doc, "Исходные данные: типичное КЛКТ — 500–600 DICOM-файлов, объём после реконструкции порядка 576×768×768 вокселей, единицы Хаунсфилда; нижняя треть лица, билатеральное расположение ВНЧС.")
    add_bullet(doc, "Почему двухэтапный подход: прямой анализ полного 3D-объёма нереалистичен (объём данных, память GPU). Этап 1 — локализация центров суставов. Этап 2 — анализ локальных областей (ROI): сегментация и/или классификация.")
    add_bullet(doc, "Требования к системе: полная автоматизация, клинически приемлемая точность, приемлемое время обработки, удобный интерфейс и поддержка DICOM.")
    add_paragraph_custom(doc, "Источники: Nir/otchet_nir_magistratura_sem1.md (раздел «Постановка задачи»).")

    add_heading_custom(doc, "3. Архитектура системы и первые шаги реализации (≈1–1,2 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Выбор архитектуры: распределённая система — тяжёлые ML-вычисления на сервере с GPU, лёгкий клиент для врача.")
    add_bullet(doc, "Три компонента: iOS-приложение (Swift/SwiftUI) — загрузка и парсинг DICOM, визуализация; Backend (Swift Vapor) — REST API, задачи, SQLite, вызов ML-сервиса; ML Service (Python, FastAPI) — препроцессинг, инференс моделей, JSON-ответы.")
    add_bullet(doc, "Что уже сделано: схема взаимодействия (REST); Backend — загрузка DICOM, создание задач, вызов ML, сохранение результатов; ML Service — пайплайн DICOM, модель детекции TMJ; iOS — MAA, парсинг DICOM, отображение координат и bbox.")
    add_bullet(doc, "Форматы данных: DICOM на входе; JSON для API (id задачи, координаты центров, bbox, при наличии — класс и уверенность).")
    add_paragraph_custom(doc, "Источники: README.md, PROJECT_PRESENTATION.md, Nir/otchet_nir_magistratura_sem1.md (Глава 1), TMJ_DETECTION_SETUP.md.")

    add_heading_custom(doc, "4. Алгоритмы машинного обучения: детекция ВНЧС (≈1–1,2 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Задача первого этапа: автоматическая локализация центров левого и правого ВНЧС в полном 3D-объёме КЛКТ (регрессия координат).")
    add_bullet(doc, "Модель детекции: 3D CNN (TMJDetectorLarge), энкодер + регрессионная «голова». Вход: даунсемплированный объём 96×128×128. Выход: координаты центров обоих суставов. Около 14,4 млн параметров.")
    add_bullet(doc, "Данные и обучение: датасет 37 аннотированных КЛКТ; аугментации (сдвиги, повороты, масштабирование, яркость, шум); потери L1/Smooth L1; Adam, ранняя остановка.")
    add_bullet(doc, "Результаты: на валидации MAE порядка 97 px (цель — снижение до <50 px или пересчёт в мм). Стабильная сходимость, детекция обоих суставов на полных 3D-снимках.")
    add_bullet(doc, "Интеграция: модель в ML Service; по координатам вырезаются ROI для этапа сегментации/классификации.")
    add_paragraph_custom(doc, "Источники: Nir/otchet_nir_magistratura_sem1.md (Глава 2), PROJECT_PRESENTATION.md, тезисы_конференция.txt, TMJ_DETECTION_SETUP.md.")

    add_heading_custom(doc, "5. Инструменты данных и планы развития (≈0,5–0,7 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Инструменты: скрипты организации датасета из DICOM; веб-инструмент разметки ROI (координаты центров ВНЧС).")
    add_bullet(doc, "Текущее состояние: 37 размеченных исследований для детекции; train/val; обучение и валидация детектора проводятся.")
    add_bullet(doc, "Планы: увеличение объёма данных; доведение точности локализации (MAE <50 px); второй этап — 3D U-Net в ROI и/или классификатор норма/патология; апробация на клинических данных; при необходимости — интерпретируемость (Grad-CAM), расширение классов.")
    add_paragraph_custom(doc, "Источники: Nir/otchet_nir_magistratura_sem1.md (Заключение), MLService/tools/, MLService/experiments/README.md.")

    add_heading_custom(doc, "6. Заключение (≈0,3–0,5 стр.)", 2)
    add_paragraph_custom(doc, "О чём писать:")
    add_bullet(doc, "Резюме: сформулирована задача анализа ВНЧС по КЛКТ; двухэтапная схема; реализована архитектура (iOS, Backend, ML Service); обучена и интегрирована модель детекции на 37 исследованиях; созданы инструменты разметки и пайплайн.")
    add_bullet(doc, "Подчеркнуть: начальный этап; система готовится к второму этапу (сегментация/классификация) и апробации в клинике.")
    add_bullet(doc, "Значимость для поддержки принятия решений в стоматологии и челюстно-лицевой диагностике.")
    add_paragraph_custom(doc, "Источники: тезисы_конференция.txt, Nir/otchet_nir_magistratura_sem1.md (Заключение).")

    add_heading_custom(doc, "Итоговая структура доклада (для проверки объёма)", 2)
    add_paragraph_custom(doc, "1. Введение и актуальность — 0,7–1 стр. (тема, актуальность, контекст, цель доклада).")
    add_paragraph_custom(doc, "2. Постановка задачи — 1–1,2 стр. (вход/выход, данные, двухэтапный подход, требования).")
    add_paragraph_custom(doc, "3. Архитектура и реализация — 1–1,2 стр. (iOS, Backend, ML Service; что сделано).")
    add_paragraph_custom(doc, "4. ML: детекция ВНЧС — 1–1,2 стр. (модель, датасет, обучение, метрики, интеграция).")
    add_paragraph_custom(doc, "5. Данные и планы — 0,5–0,7 стр. (инструменты, текущий датасет, дальнейшие шаги).")
    add_paragraph_custom(doc, "6. Заключение — 0,3–0,5 стр. (резюме, перспективы, значимость). Итого: ~5 стр.")

    add_heading_custom(doc, "Рекомендации при написании", 2)
    add_bullet(doc, "Единообразие цифр: использовать 37 аннотированных КЛКТ-исследований (в тезисах встречается 470 — уточнить у научного руководителя).")
    add_bullet(doc, "Метрики: указать MAE в пикселях (97 px) и при возможности пересчитать в мм по разрешению снимка.")
    add_bullet(doc, "Тон: подчёркивать постановку задачи, обоснование решений и достигнутые результаты; планы — короче.")
    add_bullet(doc, "Литература: 2–5 ключевых источников по образцу из тезисы_конференция.txt.")

    repo_root = Path(__file__).resolve().parent.parent
    pres_dir = repo_root / "docs" / "presentation"
    out_path = pres_dir / "план_доклада_конференция.docx"
    doc.save(str(out_path))
    print(f"Создан файл: {out_path}")

if __name__ == "__main__":
    main()
