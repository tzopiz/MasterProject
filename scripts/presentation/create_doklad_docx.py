#!/usr/bin/env python3
"""
Создаёт доклад_конференция.docx из доклад_конференция_начальный_этап.md.
Текст: Times New Roman 14 pt, заголовки: 16 pt, интервал 1,5.
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    raise SystemExit("Установите: pip install python-docx")

_REPO = Path(__file__).resolve().parent.parent
_PRES = _REPO / "docs" / "presentation"
MD_FILE = _PRES / "доклад_конференция_начальный_этап.md"
OUT_DOCX = _PRES / "доклад_конференция.docx"


def set_run_font(run, size_pt: int, bold: bool = False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_heading_custom(doc, text: str, is_title: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    p.paragraph_format.space_before = Pt(18 if is_title else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_paragraph_custom(doc, text: str):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    set_run_font(run, 14, bold=False)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p


def main():
    if not MD_FILE.exists():
        raise SystemExit(f"Не найден файл: {MD_FILE}")

    text = MD_FILE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if line.startswith("# "):
            # Заголовок первого уровня (название доклада)
            add_heading_custom(doc, stripped[2:].strip(), is_title=True)
            i += 1
            continue

        if line.startswith("## "):
            # Заголовок второго уровня
            add_heading_custom(doc, stripped[3:].strip(), is_title=False)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Собираем абзац (до пустой строки или следующего заголовка)
        para_lines = []
        while i < len(lines):
            ln = lines[i]
            if ln.strip() == "" or ln.startswith("#"):
                break
            para_lines.append(ln.strip())
            i += 1
        add_paragraph_custom(doc, " ".join(para_lines))

    doc.save(str(OUT_DOCX))
    print(f"Создан файл: {OUT_DOCX}")


if __name__ == "__main__":
    main()
